"""Convert .docx to markdown.

Paragraphs and inline tables are extracted. OLE objects and embedded images
are not extracted.
"""

from pathlib import Path


def _escape(text: str) -> str:
    return text.replace("|", "\\|").replace("\n", " ").replace("\r", "")


def _heading_level(style_name: str) -> int | None:
    name = style_name.lower()
    if name.startswith("heading"):
        parts = name.split()
        try:
            return int(parts[-1])
        except ValueError:
            return 1
    return None


def convert(src: Path, dst: Path) -> None:
    from docx import Document

    doc = Document(src)
    lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        level = _heading_level(para.style.name)
        if level is not None:
            lines.append(f"{'#' * level} {text}")
        elif text:
            lines.append(text)
        else:
            lines.append("")

    for table in doc.tables:
        lines.append("")
        header_cells = [_escape(c.text.strip()) for c in table.rows[0].cells]
        lines.append("| " + " | ".join(header_cells) + " |")
        lines.append("| " + " | ".join("---" for _ in header_cells) + " |")
        for row in table.rows[1:]:
            cells = [_escape(c.text.strip()) for c in row.cells]
            while len(cells) < len(header_cells):
                cells.append("")
            lines.append("| " + " | ".join(cells[: len(header_cells)]) + " |")
        lines.append("")

    dst.parent.mkdir(parents=True, exist_ok=True)
    dst.write_text("\n".join(lines), encoding="utf-8")
